"""Generate the ToolTune Interview Prep Guide as a Word document."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def resource(title, url, desc, rtype="Blog"):
    colors = {
        "Video": RGBColor(0xFF, 0, 0),
        "Blog": RGBColor(0, 0x66, 0xCC),
        "Course": RGBColor(0, 0x88, 0),
        "GitHub": RGBColor(0x33, 0x33, 0x33),
        "Docs": RGBColor(0xFF, 0x99, 0),
        "Tool": RGBColor(0x99, 0, 0x99),
    }
    p = doc.add_paragraph()
    tr = p.add_run(f"[{rtype}] ")
    tr.bold = True
    tr.font.color.rgb = colors.get(rtype, RGBColor(0x66, 0x66, 0x66))
    t = p.add_run(title)
    t.bold = True
    t.font.size = Pt(12)
    p2 = doc.add_paragraph()
    lr = p2.add_run(url)
    lr.font.color.rgb = RGBColor(0, 0x66, 0xCC)
    lr.font.size = Pt(10)
    lr.underline = True
    doc.add_paragraph(desc)
    doc.add_paragraph("")


def tip(text):
    p = doc.add_paragraph()
    t = p.add_run("INTERVIEW TIP: ")
    t.bold = True
    t.font.color.rgb = RGBColor(0xCC, 0x66, 0)
    p.add_run(text)
    doc.add_paragraph("")


# ====== TITLE PAGE ======
doc.add_paragraph("")
t = doc.add_heading("ToolTune Interview Prep Guide", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs:
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r.font.size = Pt(28)
s = doc.add_paragraph(
    "From Zero to Interview-Ready: LLM Fine-Tuning, GRPO & Inference"
)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = doc.add_paragraph(
    "Prepared for: Fireworks AI MTS Interview  |  Project: ToolTune"
)
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ====== TOC ======
heading("Table of Contents")
for i, item in enumerate(
    [
        "How LLMs Work (Foundations)",
        "Tokenization",
        "Fine-Tuning Basics (SFT)",
        "LoRA & QLoRA",
        "RLHF & Policy Optimization",
        "GRPO (What ToolTune Uses)",
        "Reward Design & Reward Hacking",
        "Agentic AI & Tool Use",
        "LLM Inference & Serving (Fireworks Core)",
        "Production ML",
        "5-Day Study Plan",
        "Code to Resource Mapping",
        "Interview Questions & Answers",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {item}")
doc.add_page_break()

# ====== LEVEL 1 ======
heading("Level 1: How LLMs Work (Foundations)")
doc.add_paragraph("Start here if you have zero ML background.")

resource(
    "3Blue1Brown: Neural Networks Series",
    "https://www.youtube.com/playlist?list=PLZHQObOWTQDNU6R1_67000Dx_ZCJB-3pi",
    "Gold standard visual math. Neural networks, gradient descent, backpropagation, transformers, and attention.",
    "Video",
)
resource(
    "Andrej Karpathy: Intro to Large Language Models (1hr)",
    "https://www.youtube.com/watch?v=zjkBMFhNj_g",
    "Non-technical overview of LLMs, training, fine-tuning, and the LLM OS concept. Perfect first video.",
    "Video",
)
resource(
    "Andrej Karpathy: Let's Build GPT From Scratch (2hr)",
    "https://www.youtube.com/watch?v=kCc8FmEb1nY",
    "Builds a Transformer from scratch in Python. Tokenization, self-attention, positional encoding, multi-head attention. THE single best resource for understanding transformers through code.",
    "Video",
)
resource(
    "Jay Alammar: The Illustrated Transformer",
    "https://jalammar.github.io/illustrated-transformer/",
    "Most-cited visual explainer of Transformers. Referenced at Stanford, Harvard, MIT, CMU.",
    "Blog",
)
resource(
    "Transformer Explainer (Interactive Demo)",
    "https://poloclub.github.io/transformer-explainer/",
    "Georgia Tech interactive tool - runs GPT-2 live in your browser. Watch tokens flow through attention layers.",
    "Tool",
)
resource(
    "Lilian Weng: Attention? Attention!",
    "https://lilianweng.github.io/posts/2018-06-24-attention/",
    "Deep technical survey of attention mechanisms. More thorough than Alammar. Great companion read.",
    "Blog",
)
tip(
    'If asked "explain how a transformer works": tokenization -> embeddings -> multi-head self-attention (Q/K/V) -> feed-forward layers -> softmax over vocabulary. Mention attention is O(n^2) in sequence length.'
)
doc.add_page_break()

# ====== LEVEL 2 ======
heading("Level 2: Tokenization")
resource(
    "Andrej Karpathy: Let's Build the GPT Tokenizer",
    "https://www.youtube.com/watch?v=zduSFxRajkE",
    "2-hour deep dive into BPE tokenization from scratch. Relevant because our traces use special XML tags (<think>, <tool_call>).",
    "Video",
)
resource(
    "HuggingFace: Summary of Tokenizers",
    "https://huggingface.co/docs/transformers/en/tokenizer_summary",
    "Official docs on BPE, WordPiece, and SentencePiece tokenizers. Short and practical.",
    "Docs",
)
tip(
    '"Our Qwen2.5 uses BPE with 151K vocab. We set pad_token = eos_token. max_length = 2048 tokens. Capped GRPO at 256 tokens to fit on T4."'
)
doc.add_page_break()

# ====== LEVEL 3 ======
heading("Level 3: Fine-Tuning Basics (SFT)")
doc.add_paragraph(
    "SFT is step 1 in ToolTune. Teaches format and basic tool-use by training on example traces."
)
resource(
    "HuggingFace LLM Course",
    "https://huggingface.co/learn/llm-course/chapter1/1",
    "Free course covering the HuggingFace ecosystem. Chapters 1-6 cover everything for SFT.",
    "Course",
)
resource(
    "Maxime Labonne: Beginner's Guide to LLM Fine-Tuning",
    "https://mlabonne.github.io/blog/posts/A_Beginners_Guide_to_LLM_Finetuning.html",
    "Practical walkthrough fine-tuning a 7B model. Very similar to our ToolTune SFT setup.",
    "Blog",
)
resource(
    "Maxime Labonne: LLM Course (38k+ GitHub stars)",
    "https://github.com/mlabonne/llm-course",
    "Most popular open-source LLM course. Full stack: fundamentals, SFT, DPO, RLHF, quantization, deployment.",
    "GitHub",
)
resource(
    "Sebastian Raschka: Practical Tips for Fine-Tuning LLMs",
    "https://magazine.sebastianraschka.com/p/practical-tips-for-finetuning-llms",
    "Battle-tested tips: learning rate, batch size, epochs, LoRA vs full fine-tuning, common mistakes.",
    "Blog",
)
resource(
    "HuggingFace TRL: SFTTrainer Docs",
    "https://huggingface.co/docs/trl/en/sft_trainer",
    "Official docs for the exact trainer class we use. dataset_text_field, packing, chat templates.",
    "Docs",
)
tip(
    '"SFT = supervised learning on (prompt, completion) pairs. Loss = cross-entropy on completion tokens only. 450 synthetic traces. After 1 epoch on T4 (27 min), 93% token accuracy. SFT teaches format; GRPO teaches optimization."'
)
doc.add_page_break()

# ====== LEVEL 4 ======
heading("Level 4: LoRA & QLoRA (Parameter-Efficient Fine-Tuning)")
doc.add_paragraph(
    "These let us fine-tune 7B on a single T4. Critical for Fireworks since they serve LoRA adapters for multi-tenant fine-tuning."
)
resource(
    "Sebastian Raschka: LoRA and QLoRA Explained",
    "https://sebastianraschka.com/blog/2023/llm-finetuning-lora.html",
    "Best visual + technical explainer. Low-rank decomposition (W = W0 + AB), rank r, alpha scaling, practical advice. THE one resource to read.",
    "Blog",
)
resource(
    "HuggingFace PEFT: Parameter-Efficient Fine-Tuning",
    "https://huggingface.co/blog/peft",
    "Official intro to PEFT library. LoRA, prefix tuning, prompt tuning with code examples.",
    "Blog",
)
resource(
    "HuggingFace: 4-bit Quantization with BitsAndBytes",
    "https://huggingface.co/blog/4bit-transformers-bitsandbytes",
    "NF4 quantization, double quantization, mixed-precision compute. Maps directly to our BitsAndBytesConfig.",
    "Blog",
)
resource(
    "HuggingFace PEFT Docs",
    "https://huggingface.co/docs/peft/en/index",
    "Reference for LoraConfig params (r, lora_alpha, target_modules) and PeftModel API.",
    "Docs",
)
resource(
    "HuggingFace: BitsAndBytes Quantization Docs",
    "https://huggingface.co/docs/transformers/en/quantization/bitsandbytes",
    "Official docs for all BitsAndBytesConfig quantization params we use.",
    "Docs",
)
tip(
    '"QLoRA r=16, alpha=32, targeting q/k/v/o/gate/up/down layers. Base model 4-bit NF4 (~3.5GB), adapter ~77MB. VRAM: ~10GB SFT, ~6-9GB GRPO on 15GB T4. Adapter is <1% of params but captures the full behavioral change."'
)
doc.add_page_break()

# ====== LEVEL 5 ======
heading("Level 5: RLHF & Policy Optimization")
doc.add_paragraph("RLHF made ChatGPT work. Background for GRPO.")
resource(
    "HuggingFace: Illustrating RLHF",
    "https://huggingface.co/blog/rlhf",
    "THE canonical visual guide. 3-phase pipeline: SFT, reward model training, PPO optimization.",
    "Blog",
)
resource(
    "Chip Huyen: RLHF Explained",
    "https://huyenchip.com/2023/05/02/rlhf.html",
    "In-depth: how ChatGPT-style training works, why RLHF works, relationship to hallucination.",
    "Blog",
)
resource(
    "Cameron Wolfe: PPO for LLMs",
    "https://cameronrwolfe.substack.com/p/proximal-policy-optimization-ppo",
    "RL fundamentals through PPO's clipping objective. Need this to explain why GRPO improves on PPO.",
    "Blog",
)
tip(
    '"Why GRPO over PPO? PPO needs a critic network = doubles memory. GRPO replaces it with group-relative advantages: generate G completions, normalize rewards within group. Same quality, half memory. This is what DeepSeek used for R1."'
)
doc.add_page_break()

# ====== LEVEL 6 ======
heading("Level 6: GRPO - Group Relative Policy Optimization")
doc.add_paragraph("THE most important section. GRPO is the heart of ToolTune.")
resource(
    "HuggingFace TRL: GRPOTrainer Docs",
    "https://huggingface.co/docs/trl/en/grpo_trainer",
    "Official docs for our trainer class. num_generations, beta, generation_kwargs, reward functions.",
    "Docs",
)
resource(
    "HuggingFace Blog: Open R1 (GRPO Guide)",
    "https://huggingface.co/blog/open_r1",
    "HF's open reproduction of DeepSeek R1 using GRPO. Practical setup, reward design, scaling.",
    "Blog",
)
resource(
    "Maxime Labonne: Fine-tune a Reasoning Model with GRPO",
    "https://mlabonne.github.io/blog/posts/2025-02-03_GRPO.html",
    "Hands-on GRPO tutorial with TRL. Dataset prep, reward implementation, config. Similar to ToolTune.",
    "Blog",
)

p = doc.add_paragraph()
p.add_run("GRPO Algorithm (memorize this):").bold = True
doc.add_paragraph(
    "1. Generate G completions per prompt from current policy\n"
    "2. Score each with reward function: r_1, r_2, ..., r_G\n"
    "3. Advantages: advantage_i = (r_i - mean(r)) / std(r)\n"
    "4. Loss = -advantage_i * log P(completion_i) + beta * KL(policy || reference)\n"
    "5. beta=0.04 prevents drift from SFT checkpoint"
)
tip(
    '"We use GRPO with num_generations=2, beta=0.04, composite reward (task=1.0, tools=0.3, restraint=0.1, planning=0.1, recovery=0.1, loop_penalty=-0.1). 60 steps, max_new_tokens=256 on T4."'
)
doc.add_page_break()

# ====== LEVEL 7 ======
heading("Level 7: Reward Design & Reward Hacking")
resource(
    "Lilian Weng: Reward Hacking in RL",
    "https://lilianweng.github.io/posts/2024-11-28-reward-hacking/",
    "How agents exploit reward flaws. Goodhart's Law, overoptimization, LLM alignment examples, mitigations.",
    "Blog",
)
resource(
    "Anthropic: Measuring Faithfulness in Chain-of-Thought",
    "https://www.anthropic.com/research/measuring-faithfulness-in-chain-of-thought-reasoning",
    "Do LLMs actually use their reasoning or produce plausible post-hoc? Relevant to our <think> blocks.",
    "Blog",
)
p = doc.add_paragraph()
p.add_run("Our ToolTune Reward Function:").bold = True
doc.add_paragraph(
    "task_completion  = 1.0  (correct answer?)\n"
    "tool_accuracy    = 0.3  (right tools, valid args?)\n"
    "restraint        = 0.1  (avoided tools when unnecessary?)\n"
    "planning         = 0.1  (used <think> blocks?)\n"
    "error_recovery   = 0.1  (recovered from tool errors?)\n"
    "loop_penalty     = -0.1 (per excess tool call)"
)
tip(
    '"Task completion at 1.0 dominates - prevents reward hacking where model produces perfect tool calls but wrong answers. Restraint (0.1) teaches GRPO to answer 2+2 directly instead of calling calculator."'
)
doc.add_page_break()

# ====== LEVEL 8 ======
heading("Level 8: Agentic AI & Tool Use")
resource(
    "Lilian Weng: LLM Powered Autonomous Agents",
    "https://lilianweng.github.io/posts/2023-06-23-agent/",
    "Most comprehensive blog on LLM agents. Planning, memory, tool use, ReAct pattern. Maps to our agentic_loop.py.",
    "Blog",
)
resource(
    "HuggingFace: Building Agents Tutorial",
    "https://huggingface.co/learn/cookbook/en/agents",
    "Hands-on tutorial building tool-using agents. Similar architecture to ToolTune.",
    "Course",
)
p = doc.add_paragraph()
p.add_run("Our ReAct Loop (from train/agentic_loop.py):").bold = True
doc.add_paragraph(
    "1. Model gets system prompt with tool JSON schemas\n"
    '2. Outputs <think>reasoning</think>\n'
    '3. Outputs <tool_call>{"name":"calculator","arguments":{"expression":"347*892"}}</tool_call>\n'
    "4. System executes tool, returns <observation>309524</observation>\n"
    "5. Model reads observation, decides: answer or another tool call\n"
    "6. Outputs <answer>309524</answer>\n"
    "\nMax 5 steps. Loops get penalized."
)
tip(
    '"Key insight: during GRPO training, the model hallucinates its own observations - no real tools in the loop. Reward scores final answer vs ground truth. Real tool execution only at inference."'
)
doc.add_page_break()

# ====== LEVEL 9 ======
heading("Level 9: LLM Inference & Serving (FIREWORKS AI CORE)")
p = doc.add_paragraph()
r = p.add_run("THIS IS THE MOST IMPORTANT SECTION FOR YOUR INTERVIEW")
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0, 0)

resource(
    "vLLM Blog: PagedAttention Explained",
    "https://blog.vllm.ai/2023/06/20/vllm.html",
    "Original vLLM blog: OS-style virtual memory for KV cache. Eliminates fragmentation. Core concept for Fireworks.",
    "Blog",
)
resource(
    "Anyscale: Continuous Batching for LLM Inference",
    "https://www.anyscale.com/blog/continuous-batching-llm-inference",
    "Why naive batching wastes GPU. Continuous batching = iteration-level scheduling. Core to Fireworks throughput.",
    "Blog",
)
resource(
    "Fireworks AI Engineering Blog",
    "https://fireworks.ai/blog",
    "READ THEIR OWN BLOG. Shows you did homework. Their optimizations, architecture, product decisions.",
    "Blog",
)
resource(
    "Fireworks: FireAttention (4x faster than vLLM)",
    "https://fireworks.ai/blog/fire-attention-serving-open-source-models-4x-faster-than-vllm-by-quantizing-with-no-tradeoffs",
    "Their custom CUDA kernel for quantized attention. Understanding this shows deep interview prep.",
    "Blog",
)
resource(
    "Jay Alammar: The Illustrated GPT-2 (Inference)",
    "https://jalammar.github.io/illustrated-gpt2/",
    "Visual walkthrough of autoregressive inference token by token. Understand why KV caching matters.",
    "Blog",
)

p = doc.add_paragraph()
p.add_run("Key Inference Concepts to Memorize:").bold = True
for concept in [
    "KV CACHE: Store Key/Value matrices from previous tokens so each new token only computes its own attention. Without: O(n^2). With: O(n). Memory = seq_len * layers * hidden * 2.",
    "PAGED ATTENTION: Divide KV cache into fixed pages + page table (like OS virtual memory). No fragmentation. Enables sharing pages for common prefixes.",
    "CONTINUOUS BATCHING: Insert/remove requests each iteration. Short requests free GPU immediately.",
    "SPECULATIVE DECODING: Small draft model predicts N tokens, large model verifies all N in one pass. 2-3x speedup.",
    "LoRA MULTI-TENANT: One base model + swap ~80MB adapters per request. 100+ fine-tunes on one GPU.",
    "QUANTIZATION FOR SERVING: 4-bit/8-bit inference = 2-4x memory savings via INT8 tensor cores.",
]:
    doc.add_paragraph(concept, style="List Bullet")

tip(
    '"I built the training pipeline that produces LoRA adapters. Those adapters are exactly what Fireworks serves - small weights swapped on quantized base models. I understand both training AND serving: low latency, high throughput, KV cache management, multi-tenant LoRA serving."'
)
doc.add_page_break()

# ====== LEVEL 10 ======
heading("Level 10: Production ML & Developer Experience")
resource(
    "FastAPI Official Tutorial",
    "https://fastapi.tiangolo.com/tutorial/",
    "Our playground uses FastAPI. Pydantic validation, SSE streaming, async endpoints.",
    "Docs",
)
resource(
    "Chip Huyen: ML Systems Blog",
    "https://huyenchip.com/blog/",
    "Real-world ML systems: evaluation, monitoring, deployment. Essential for MTS-level thinking.",
    "Blog",
)
doc.add_page_break()

# ====== STUDY PLAN ======
heading("5-Day Study Plan")
days = [
    (
        "Day 1: Foundations + SFT (2.5 hrs)",
        [
            'Watch: Karpathy "Intro to LLMs" (1 hr)',
            "Read: Jay Alammar Illustrated Transformer (30 min)",
            "Read: HuggingFace RLHF blog (20 min)",
            "Read: Raschka LoRA blog (30 min)",
            "Play: Transformer Explainer demo (15 min)",
        ],
    ),
    (
        "Day 2: LoRA + QLoRA Deep Dive (2 hrs)",
        [
            "Read: HF PEFT blog + 4-bit BnB blog (40 min)",
            "Read: Maxime Labonne fine-tuning guide (30 min)",
            "Review: YOUR train/sft_tooltune.py code (30 min)",
            "Read: TRL SFTTrainer docs (20 min)",
        ],
    ),
    (
        "Day 3: RLHF + GRPO - MOST IMPORTANT (2.5 hrs)",
        [
            "Read: Chip Huyen RLHF blog (30 min)",
            "Read: Cameron Wolfe PPO blog (45 min)",
            "Read: TRL GRPOTrainer docs + Labonne GRPO tutorial (50 min)",
            "Review: YOUR train/grpo_tooltune.py + reward.py (30 min)",
        ],
    ),
    (
        "Day 4: Inference & Fireworks - CRITICAL (2.5 hrs)",
        [
            "Read: vLLM PagedAttention blog (30 min)",
            "Read: Anyscale continuous batching blog (20 min)",
            "Read: ALL Fireworks AI blog posts (45 min)",
            "Read: Fireworks FireAttention blog (20 min)",
            "Read: Lilian Weng reward hacking blog (30 min)",
        ],
    ),
    (
        "Day 5: Project Review + Mock Interview (2 hrs)",
        [
            "Walk through every train/ file (30 min)",
            "Walk through agentic_loop.py + tools/ (20 min)",
            "Practice Q&A below OUT LOUD (45 min)",
            "Walk through playground/ code (15 min)",
        ],
    ),
]
for title, items in days:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")
doc.add_page_break()

# ====== CODE MAPPING ======
heading("Code to Resource Mapping")
for f, c in [
    ("train/sft_tooltune.py", "SFTTrainer, QLoRA, BitsAndBytes, LoRA, bf16 vs fp16"),
    ("train/grpo_tooltune.py", "GRPOTrainer, KL penalty (beta), generation config, reward integration"),
    ("train/reward.py", "Composite reward design, answer matching, tool validation, loop penalty"),
    ("train/agentic_loop.py", "ReAct pattern, system prompt + tool schemas, regex parsing, SSE events"),
    ("tools/registry.py", "Tool dispatch, parameter validation, error injection, JSON schema"),
    ("tools/calculator.py", "AST-based safe eval (no code injection)"),
    ("playground/app.py + routes.py", "FastAPI, CORS, SSE streaming, async endpoints"),
    ("playground/agent.py", "vLLM client, live vs demo inference, trace verification"),
]:
    p = doc.add_paragraph()
    r = p.add_run(f)
    r.bold = True
    r.font.name = "Consolas"
    doc.add_paragraph(c)
doc.add_page_break()

# ====== Q&A ======
heading("Interview Questions & Answers")
for q, a in [
    (
        "Walk me through ToolTune end to end.",
        "Problem: base LLMs are bad at tool use. Pipeline: (1) Generate 500 tasks across 4 tiers, (2) SFT warmup on 450 synthetic traces for ReAct format, (3) GRPO with composite reward to optimize behavior. QLoRA (4-bit NF4, r=16) on Qwen2.5-7B-Instruct, single T4 GPU. Result: measurable improvement base -> SFT -> GRPO on task completion, tool accuracy, and restraint. FastAPI playground demos traces side-by-side.",
    ),
    (
        "What is GRPO and why over PPO?",
        "GRPO generates G completions per prompt, normalizes rewards within group for advantages. No critic network = ~50% less memory. Chose it because: (1) T4 memory tight, (2) reward is programmatic/verifiable, (3) SOTA post-DeepSeek R1.",
    ),
    (
        "How did you design the reward function?",
        "Composite with 6 components weighted by importance: task completion (1.0) dominant since correct answers matter most. Tool accuracy (0.3) checks precision/recall. Restraint (0.1) rewards NOT calling tools unnecessarily. Planning (0.1) checks <think> blocks. Recovery (0.1). Loop penalty (-0.1/excess). Asymmetric weighting prevents reward hacking.",
    ),
    (
        "What is QLoRA and why for serving?",
        "4-bit base + low-rank adapters. Training: 7B in ~6-10GB. Serving (Fireworks): one base in GPU memory, swap ~80MB adapters per request. Multi-tenant: hundreds of fine-tunes, one GPU. Adapter <1% of params.",
    ),
    (
        "Explain KV caching.",
        "Store K,V from previous tokens so each new token only computes its own attention. Without: O(n^2). With: O(n). Memory grows linearly. PagedAttention manages with virtual memory paging to eliminate fragmentation.",
    ),
    (
        "How does Fireworks serve models fast?",
        "(1) PagedAttention for KV cache management, (2) Continuous batching for GPU utilization, (3) FireAttention custom CUDA kernels for quantized inference, (4) LoRA adapter swapping for multi-tenant serving.",
    ),
    (
        "What would you improve about ToolTune?",
        "(1) Real tool execution during GRPO training (environment-in-the-loop), (2) More GRPO steps on A100 for stronger behavioral separation, (3) Live vLLM endpoint serving all 3 variants with <100ms TTFT, mapping to Fireworks production.",
    ),
]:
    p = doc.add_paragraph()
    qr = p.add_run(f"Q: {q}")
    qr.bold = True
    qr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p2 = doc.add_paragraph()
    ar = p2.add_run(f"A: {a}")
    ar.font.size = Pt(10)
    ar.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph("")

# ====== SAVE ======
downloads = os.path.join(os.path.expanduser("~"), "Downloads")
filepath = os.path.join(downloads, "ToolTune_Interview_Prep_Guide.docx")
doc.save(filepath)
print(f"SAVED: {filepath}")
